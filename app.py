import streamlit as st
import uuid
from datetime import datetime, timezone
from pathlib import Path
import csv
from io import BytesIO
import os

# ---------- 啟動時清快取，降低雲端殘留影響 ----------
try:
    st.cache_data.clear()
    st.cache_resource.clear()
except Exception:
    pass

# ---------- 基本設定 ----------
st.set_page_config(page_title="影響力平台", page_icon="✨", layout="wide")

DATA_DIR = Path("data")
DATA_DIR.mkdir(parents=True, exist_ok=True)
CASES_CSV = DATA_DIR / "cases.csv"

# 可選套件（不裝也能跑）
try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None
    Pt = None

# ---------- 管理密鑰（secrets > env > demo） ----------
ADMIN_KEY = None
try:
    ADMIN_KEY = st.secrets.get("ADMIN_KEY", None)
except Exception:
    ADMIN_KEY = None
if not ADMIN_KEY:
    ADMIN_KEY = os.environ.get("ADMIN_KEY", "demo")  # 本機/未設定時可用 demo

# ---------- 狀態 ----------
if "page" not in st.session_state:
    st.session_state.page = "首頁"
if "last_case_id" not in st.session_state:
    st.session_state.last_case_id = ""
if "admin_ok" not in st.session_state:
    st.session_state.admin_ok = False

# ---------- 共用：頁尾 ----------
def footer():
    st.markdown("---")
    st.markdown(
        """
        <div style='display: flex; justify-content: center; align-items: center; gap: 1.5em; font-size: 14px; color: gray;'>
          <a href='?' style='color:#006666; text-decoration: underline;'>《影響力》傳承策略平台</a>
          <a href='https://gracefo.com' target='_blank'>永傳家族辦公室</a>
          <a href='mailto:123@gracefo.com'>123@gracefo.com</a>
        </div>
        """,
        unsafe_allow_html=True
    )

# ---------- CSV 儲存/讀取 ----------
CSV_HEADERS = [
    "ts","case_id","name","mobile","email","marital","children","special",
    "equity","real_estate","financial","insurance_cov",
    "focus","total_assets","liq_low","liq_high","gap_low","gap_high"
]

def append_case_row(row: dict):
    need_header = not CASES_CSV.exists()
    with CASES_CSV.open("a", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=CSV_HEADERS)
        if need_header:
            w.writeheader()
        w.writerow({k: row.get(k, "") for k in CSV_HEADERS})

@st.cache_data(show_spinner=False)
def load_all_cases():
    """回傳 (rows, df or None)。"""
    if not CASES_CSV.exists():
        return [], (pd.DataFrame(columns=CSV_HEADERS) if pd else None)
    rows = []
    with CASES_CSV.open("r", newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        for row in r:
            rows.append(row)
    if pd:
        df = pd.DataFrame(rows)
        # 數字轉型
        for col in ["children","equity","real_estate","financial","insurance_cov",
                    "total_assets","liq_low","liq_high","gap_low","gap_high"]:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")
        # 時間排序
        if "ts" in df.columns:
            df["ts"] = pd.to_datetime(df["ts"], errors="coerce")
            df = df.sort_values("ts", ascending=False)
        return rows, df
    return rows, None

def load_case(case_id: str):
    if not CASES_CSV.exists():
        return None
    with CASES_CSV.open("r", newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        last = None
        for row in r:
            if row.get("case_id") == case_id:
                last = row
    if not last:
        return None
    for k in ["children","equity","real_estate","financial","insurance_cov",
              "total_assets","liq_low","liq_high","gap_low","gap_high"]:
        if last.get(k):
            try:
                last[k] = int(float(last[k]))
            except Exception:
                pass
    if last.get("focus"):
        last["focus"] = [x for x in last["focus"].split("|") if x]
    else:
        last["focus"] = []
    return last

# ---------- 報告輸出 ----------
def build_docx_bytes(case_id: str, case: dict) -> bytes | None:
    if Document is None:
        return None
    doc = Document()
    if Pt is not None:
        try:
            styles = doc.styles['Normal']
            styles.font.name = 'Microsoft JhengHei'
            styles.font.size = Pt(11)
        except Exception:
            pass
    doc.add_heading("影響力平台｜傳承規劃簡版報告", level=1)
    doc.add_paragraph(f"個案編號：{case_id}")
    doc.add_paragraph(f"建立時間：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_heading("一、基本資料", level=2)
    doc.add_paragraph(f"申請人：{case.get('name') or '（未填）'}")
    doc.add_paragraph(f"婚姻：{case.get('marital','')}　子女：{case.get('children','')}")
    doc.add_paragraph(f"特殊照顧對象：{case.get('special','')}")
    doc.add_heading("二、資產概況（估）", level=2)
    doc.add_paragraph(f"- 公司股權：{case['equity']:,} 萬")
    doc.add_paragraph(f"- 不動產：{case['real_estate']:,} 萬")
    doc.add_paragraph(f"- 金融資產：{case['financial']:,} 萬")
    doc.add_paragraph(f"- 既有保單保額：{case['insurance_cov']:,} 萬")
    doc.add_paragraph(f"- 資產總額（估）：{case['total_assets']:,} 萬")
    doc.add_heading("三、交棒流動性與保障缺口（示意）", level=2)
    doc.add_paragraph(f"- 交棒流動性需求（估）：{case['liq_low']:,} – {case['liq_high']:,} 萬")
    doc.add_paragraph(f"- 可能的保障缺口：{case['gap_low']:,} – {case['gap_high']:,} 萬")
    doc.add_heading("四、您的重點關注", level=2)
    if case.get("focus"):
        for f in case["focus"]:
            doc.add_paragraph(f"• {f}")
    else:
        doc.add_paragraph("（未填）")
    doc.add_heading("五、初步建議（草案）", level=2)
    for b in [
        "以保單建立緊急流動性池，避免交棒時資金壓力。",
        "評估是否需要信託來管理特殊照顧對象或特定資產的分配節奏。",
        "針對股權與不動產，規劃適當的傳承順序與治理安排。",
        "視需要規劃遺囑，確保意願清楚、減少爭議。",
    ]:
        doc.add_paragraph(f"• {b}")
    doc.add_paragraph("\n免責聲明：本報告為初步示意，實際方案須由專業顧問複核並依相關法令辦理。")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def build_txt_bytes(case_id: str, case: dict) -> bytes:
    lines = []
    lines.append("影響力平台｜傳承規劃簡版報告")
    lines.append(f"個案編號：{case_id}")
    lines.append(f"建立時間：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    lines.append("")
    lines.append("一、基本資料")
    lines.append(f"- 申請人：{case.get('name') or '（未填）'}")
    lines.append(f"- 婚姻：{case.get('marital','')}　子女：{case.get('children','')}")
    lines.append(f"- 特殊照顧對象：{case.get('special','')}")
    lines.append("")
    lines.append("二、資產概況（估）")
    lines.append(f"- 公司股權：{case['equity']:,} 萬")
    lines.append(f"- 不動產：{case['real_estate']:,} 萬")
    lines.append(f"- 金融資產：{case['financial']:,} 萬")
    lines.append(f"- 既有保單保額：{case['insurance_cov']:,} 萬")
    lines.append(f"- 資產總額（估）：{case['total_assets']:,} 萬")
    lines.append("")
    lines.append("三、交棒流動性與保障缺口（示意）")
    lines.append(f"- 交棒流動性需求（估）：{case['liq_low']:,} – {case['liq_high']:,} 萬")
    lines.append(f"- 可能的保障缺口：{case['gap_low']:,} – {case['gap_high']:,} 萬")
    lines.append("")
    lines.append("四、您的重點關注")
    if case.get("focus"):
        lines += [f"• {f}" for f in case["focus"]]
    else:
        lines.append("（未填）")
    lines.append("")
    lines.append("五、初步建議（草案）")
    lines += [
        "• 以保單建立緊急流動性池，避免交棒時資金壓力。",
        "• 評估是否需要信託來管理特殊照顧對象或特定資產的分配節奏。",
        "• 針對股權與不動產，規劃適當的傳承順序與治理安排。",
        "• 視需要規劃遺囑，確保意願清楚、減少爭議。",
    ]
    lines.append("")
    lines.append("免責聲明：本報告為初步示意，實際方案須由專業顧問複核並依相關法令辦理。")
    return ("\n".join(lines)).encode("utf-8")

# ---------- 頁面們 ----------
def page_home():
    st.title("傳承您的影響力")
    st.write("AI 智慧 + 專業顧問，打造專屬的可視化傳承方案，確保財富與愛同時流傳。")

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("### 家族資產地圖\n將股權、不動產、保單、金融資產一次整理")
    with c2:
        st.markdown("### AI 傳承策略\n根據家族偏好與資料生成個人化方案")
    with c3:
        st.markdown("### 行動計劃表\n明確列出下一步與時間表，陪伴落地")

    st.divider()
    st.subheader("立即行動")
    a, b = st.columns(2)
    with a:
        if st.button("開始規劃（免費）", use_container_width=True):
            st.session_state.page = "診斷"
            st.rerun()
    with b:
        if st.button("預約 30 分鐘諮詢", use_container_width=True):
            st.session_state.page = "預約"
            st.rerun()

    st.caption("免責：本平台提供之計算與建議僅供初步規劃參考，請依專業顧問複核與相關法令為準。")
    footer()

def page_diagnostic():
    st.title("傳承規劃｜快速診斷（MVP）")
    st.write("填寫 60 秒，取得初步風險指標與行動建議。")

    with st.form("diag"):
        st.subheader("家庭結構")
        c1, c2, c3 = st.columns(3)
        marital = c1.selectbox("婚姻狀態", ["未婚", "已婚", "離異", "喪偶"])
        children = c2.number_input("子女數", min_value=0, max_value=10, step=1, value=2)
        special = c3.selectbox("是否有特殊照顧對象", ["否", "是"])

        st.subheader("資產概況（估算即可，單位：萬元）")
        c4, c5 = st.columns(2)
        equity = c4.number_input("公司股權估值", min_value=0, step=100, value=5000)
        real_estate = c5.number_input("不動產估值", min_value=0, step=100, value=8000)
        c6, c7 = st.columns(2)
        financial = c6.number_input("金融資產估值", min_value=0, step=100, value=3000)
        insurance_cov = c7.number_input("既有保單保額", min_value=0, step=100, value=2000)

        st.subheader("您的重點關注（可多選）")
        focus = st.multiselect(
            "選擇重點",
            ["稅務負擔", "現金流穩定", "交棒安排", "家族和諧", "跨境安排"],
            default=["現金流穩定", "交棒安排"]
        )

        st.subheader("聯絡方式")
        c8, c9 = st.columns(2)
        name = c8.text_input("姓名")
        mobile = c9.text_input("手機")
        email = st.text_input("Email")

        submitted = st.form_submit_button("產生診斷結果與 CaseID")

    if submitted:
        case_id = "YC-" + uuid.uuid4().hex[:6].upper()
        total_assets = equity + real_estate + financial
        liq_low = round(total_assets * 0.10)
        liq_high = round(total_assets * 0.20)
        gap_low = max(0, liq_low - insurance_cov)
        gap_high = max(0, liq_high - insurance_cov)

        row = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "case_id": case_id,
            "name": name, "mobile": mobile, "email": email,
            "marital": marital, "children": int(children), "special": special,
            "equity": int(equity), "real_estate": int(real_estate), "financial": int(financial),
            "insurance_cov": int(insurance_cov),
            "focus": "|".join(focus),
            "total_assets": int(total_assets),
            "liq_low": int(liq_low), "liq_high": int(liq_high),
            "gap_low": int(gap_low), "gap_high": int(gap_high),
        }
        append_case_row(row)

        st.session_state.last_case_id = case_id
        st.session_state.page = "結果"
        st.rerun()

    footer()

def page_result():
    st.title("診斷結果（簡版）")
    case_id = st.text_input("輸入 CaseID 查詢", value=st.session_state.get("last_case_id", ""))
    if not case_id:
        st.info("尚無資料。請先完成『快速診斷』產生 CaseID。")
        footer(); return

    case = load_case(case_id)
    if not case:
        st.warning("查無此 CaseID，請確認輸入是否正確。")
        footer(); return

    st.markdown(f"**個案編號：** `{case_id}`  \n**申請人：** {case.get('name') or '（未填）'}")
    st.divider()

    st.subheader("一、風險重點")
    st.write(f"- 資產總額（估）：**{case['total_assets']:,} 萬**")
    st.write(f"- 交棒流動性需求（估）：**{case['liq_low']:,}–{case['liq_high']:,} 萬**")
    st.write(f"- 現有保單保額：**{case['insurance_cov']:,} 萬**")
    st.write(f"- 可能的保障缺口範圍：**{case['gap_low']:,}–{case['gap_high']:,} 萬**")
    st.caption("說明：以上為示意試算，實際仍需依照家庭目標、法規與細部資產結構調整。")

    st.subheader("二、動作與下載")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("預約 30 分鐘諮詢"):
            st.session_state.page = "預約"
            st.rerun()
    with c2:
        docx_bytes = build_docx_bytes(case_id, case)
        if docx_bytes:
            st.download_button(
                label="下載簡版報告（.docx）",
                data=docx_bytes,
                file_name=f"{case_id}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            txt_bytes = build_txt_bytes(case_id, case)
            st.download_button(
                label="下載簡版報告（.txt）",
                data=txt_bytes,
                file_name=f"{case_id}_report.txt",
                mime="text/plain",
                use_container_width=True
            )

    footer()

# ---- 管理：案件總表（即時驗證；成功自動刷新） ----
def page_cases():
    st.title("案件總表（管理）")

    def _verify_from_sidebar():
        key_in = st.session_state.get("admin_key_sidebar", "")
        st.session_state.admin_ok = (key_in == ADMIN_KEY)
        if st.session_state.admin_ok:
            st.toast("已通過驗證，正在載入資料…", icon="✅")
            st.rerun()
        else:
            st.toast("密鑰錯誤", icon="❌")

    def _verify_from_inline():
        key_in = st.session_state.get("admin_key_inline", "")
        st.session_state.admin_ok = (key_in == ADMIN_KEY)
        if st.session_state.admin_ok:
            st.toast("已通過驗證，正在載入資料…", icon="✅")
            st.rerun()
        else:
            st.toast("密鑰錯誤", icon="❌")

    # 側邊欄驗證（on_change：按 Enter 立即驗證）
    with st.sidebar:
        st.subheader("管理密鑰")
        st.text_input("請輸入管理密鑰", type="password", key="admin_key_sidebar", on_change=_verify_from_sidebar)

    # 未通過：頁內顯示另一個即時驗證入口
    if not st.session_state.admin_ok:
        st.warning("此頁需管理密鑰。如未設定 secrets，可先用測試密鑰：`demo`。")
        st.text_input("在此輸入管理密鑰（頁內）", type="password", key="admin_key_inline", on_change=_verify_from_inline)
        footer(); return

    # 已通過：顯示總表
    rows, df = load_all_cases()
    if not rows:
        st.info("目前尚無個案資料。請先到『診斷』頁建立個案。")
        footer(); return

    st.subheader("資料檢視與篩選")
    c1, c2, _ = st.columns([2,2,1])
    kw = c1.text_input("姓名 / Email 關鍵字", "")
    recent_days = c2.number_input("僅看最近 N 天（0 表示不限制）", min_value=0, max_value=3650, value=0, step=1)

    if df is not None:
        df_view = df.copy()
        if kw:
            mask = (df_view["name"].fillna("").str.contains(kw, case=False)) | \
                   (df_view["email"].fillna("").str.contains(kw, case=False))
            df_view = df_view[mask]
        if recent_days and "ts" in df_view.columns:
            import pandas as pd
            cutoff = datetime.now(timezone.utc) - pd.Timedelta(days=recent_days)
            df_view = df_view[df_view["ts"] >= cutoff]

        st.dataframe(
            df_view[["ts","case_id","name","email","mobile","total_assets","liq_low","liq_high","gap_low","gap_high"]],
            use_container_width=True, height=420
        )
        st.markdown("—")
        selected_case = st.selectbox("選擇 CaseID 前往檢視結果", df_view["case_id"].astype(str).tolist())
    else:
        view = rows
        if kw:
            view = [r for r in view if kw.lower() in (r.get("name","") + r.get("email","")).lower()]
        if recent_days:
            cutoff = datetime.now(timezone.utc)
            def within_days(ts):
                try:
                    dt = datetime.fromisoformat(ts.replace("Z",""))
                    return (cutoff - dt).days <= recent_days
                except Exception:
                    return True
            view = [r for r in view if within_days(r.get("ts",""))]
        st.dataframe(view, use_container_width=True, height=420)
        selected_case = st.selectbox("選擇 CaseID 前往檢視結果", [r["case_id"] for r in view])

    cA, cB = st.columns(2)
    with cA:
        if st.button("前往結果頁"):
            st.session_state.last_case_id = selected_case
            st.session_state.page = "結果"
            st.rerun()
    with cB:
        with CASES_CSV.open("rb") as f:
            st.download_button(
                "下載 cases.csv", data=f.read(), file_name="cases.csv",
                mime="text/csv", use_container_width=True
            )

    footer()

def page_book():
    st.title("預約 30 分鐘線上會談")
    st.info("（正式版可嵌入 Calendly / Google 日曆 iframe）")
    with st.form("book"):
        name = st.text_input("姓名")
        phone = st.text_input("手機")
        email = st.text_input("Email")
        notes = st.text_area("想先告訴我們的情況（選填）")
        if st.form_submit_button("送出預約申請"):
            st.success("已收到預約申請，我們將盡快與您聯繫。")
    footer()

def page_advisors():
    st.title("顧問專區（MVP）")
    st.write("先上線最小功能，驗證註冊意願。正式版會加入白標報告與授權方案。")
    with st.form("adv_signup"):
        name = st.text_input("姓名 / 公司名")
        email = st.text_input("Email")
        phone = st.text_input("手機")
        brand = st.text_input("希望顯示於報告的顧問品牌名稱")
        if st.form_submit_button("建立帳號（示意）"):
            st.success("註冊成功（示意）。正式版將儲存資料並寄出歡迎信。")
    footer()

def page_plans():
    st.title("授權與高階會員方案（合規）")
    c1, c2 = st.columns(2)
    with c1:
        st.header("Starter（授權）")
        st.markdown(
            "- AI 診斷、提案草案、案例庫（基礎）  \n"
            "- 簡版報告（白標）  \n"
            "- 客戶個案版本歷程  \n"
            "- **NT$ 3,600 / 月** 或 **NT$ 36,000 / 年**"
        )
    with c2:
        st.header("Pro（高階會員）")
        st.markdown(
            "- 進階策略模組與圖像化報告客製  \n"
            "- 專屬培訓與話術庫、實戰案例包  \n"
            "- 專案協作席次 3 位（團隊版）  \n"
            "- **NT$ 12,000 / 月** 或 **NT$ 120,000 / 年**"
        )
    st.caption("合規：平台僅收授權與專業服務費，不參與佣金分配或分潤。")
    footer()

def page_privacy():
    st.title("隱私與免責聲明")
    st.write(
        "- 我們重視您的個人資料保護，僅在提供服務之目的範圍內蒐集與使用。  \n"
        "- 您可要求查詢、更正或刪除個人資料，詳情請與我們聯繫。  \n"
        "- 本平台之計算結果與建議僅供初步規劃參考，實際仍需依照家庭目標、法規與細部資產結構調整。"
    )
    footer()

# ---------- 側邊欄導航 ----------
st.sidebar.header("功能選單")
page = st.sidebar.radio(
    "選擇頁面",
    ("首頁", "診斷", "結果", "案件總表（管理）", "預約", "顧問", "方案", "隱私"),
    index=("首頁","診斷","結果","案件總表（管理）","預約","顧問","方案","隱私").index(st.session_state.page)
)
st.session_state.page = page

# ---------- 路由 ----------
if page == "首頁":
    page_home()
elif page == "診斷":
    page_diagnostic()
elif page == "結果":
    page_result()
elif page == "案件總表（管理）":
    page_cases()
elif page == "預約":
    page_book()
elif page == "顧問":
    page_advisors()
elif page == "方案":
    page_plans()
else:
    page_privacy()
