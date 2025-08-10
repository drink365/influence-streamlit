from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None
    Pt = None

from datetime import datetime, timezone

def build_docx(case_id: str, case: dict):
    if Document is None:
        return None
    doc = Document()
    if Pt is not None:
        try:
            styles = doc.styles['Normal']
            styles.font.name = 'Microsoft JhengHei'
            styles.font.size = Pt(11)
        except Exception:
            pass
    doc.add_heading("影響力平台｜傳承規劃簡版報告", level=1)
    doc.add_paragraph(f"個案編號：{case_id}")
    doc.add_paragraph(f"建立時間：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_heading("一、基本資料", level=2)
    doc.add_paragraph(f"申請人：{case.get('name') or '（未填）'}")
    doc.add_paragraph(f"婚姻：{case.get('marital','')}　子女：{case.get('children','')}")
    doc.add_paragraph(f"特殊照顧對象：{case.get('special','')}")
    doc.add_heading("二、資產概況（估）", level=2)
    doc.add_paragraph(f"- 公司股權：{case['equity']:,} 萬")
    doc.add_paragraph(f"- 不動產：{case['real_estate']:,} 萬")
    doc.add_paragraph(f"- 金融資產：{case['financial']:,} 萬")
    doc.add_paragraph(f"- 既有保單保額：{case['insurance_cov']:,} 萬")
    doc.add_paragraph(f"- 資產總額（估）：{case['total_assets']:,} 萬")
    doc.add_heading("三、交棒流動性與保障缺口（示意）", level=2)
    doc.add_paragraph(f"- 交棒流動性需求（估）：{case['liq_low']:,} – {case['liq_high']:,} 萬")
    doc.add_paragraph(f"- 可能的保障缺口：{case['gap_low']:,} – {case['gap_high']:,} 萬")
    doc.add_heading("四、您的重點關注", level=2)
    if case.get("focus"):
        for f in case["focus"]:
            doc.add_paragraph(f"• {f}")
    else:
        doc.add_paragraph("（未填）")
    doc.add_heading("五、初步建議（草案）", level=2)
    for b in [
        "以保單建立緊急流動性池，避免交棒時資金壓力。",
        "評估是否需要信託來管理特殊照顧對象或特定資產的分配節奏。",
        "針對股權與不動產，規劃適當的傳承順序與治理安排。",
        "視需要規劃遺囑，確保意願清楚、減少爭議。",
    ]:
        doc.add_paragraph(f"• {b}")
    doc.add_paragraph("\n免責聲明：本報告為初步示意，實際方案須由專業顧問複核並依相關法令辦理。")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_txt(case_id: str, case: dict) -> bytes:
    lines = []
    lines.append("影響力平台｜傳承規劃簡版報告")
    lines.append(f"個案編號：{case_id}")
    lines.append(f"建立時間：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    lines.append("")
    lines.append("一、基本資料")
    lines.append(f"- 申請人：{case.get('name') or '（未填）'}")
    lines.append(f"- 婚姻：{case.get('marital','')}　子女：{case.get('children','')}")
    lines.append(f"- 特殊照顧對象：{case.get('special','')}")
    lines.append("")
    lines.append("二、資產概況（估）")
    lines.append(f"- 公司股權：{case['equity']:,} 萬")
    lines.append(f"- 不動產：{case['real_estate']:,} 萬")
    lines.append(f"- 金融資產：{case['financial']:,} 萬")
    lines.append(f"- 既有保單保額：{case['insurance_cov']:,} 萬")
    lines.append(f"- 資產總額（估）：{case['total_assets']:,} 萬")
    lines.append("")
    lines.append("三、交棒流動性與保障缺口（示意）")
    lines.append(f"- 交棒流動性需求（估）：{case['liq_low']:,} – {case['liq_high']:,} 萬")
    lines.append(f"- 可能的保障缺口：{case['gap_low']:,} – {case['gap_high']:,} 萬")
    lines.append("")
    lines.append("四、您的重點關注")
    if case.get("focus"):
        lines += [f"• {f}" for f in case["focus"]]
    else:
        lines.append("（未填）")
    lines.append("")
    lines.append("五、初步建議（草案）")
    lines += [
        "• 以保單建立緊急流動性池，避免交棒時資金壓力。",
        "• 評估是否需要信託來管理特殊照顧對象或特定資產的分配節奏。",
        "• 針對股權與不動產，規劃適當的傳承順序與治理安排。",
        "• 視需要規劃遺囑，確保意願清楚、減少爭議。",
    ]
    lines.append("")
    lines.append("免責聲明：本報告為初步示意，實際方案須由專業顧問複核並依相關法令辦理。")
    return ("\n".join(lines)).encode("utf-8")
